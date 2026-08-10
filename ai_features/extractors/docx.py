import logging

logger = logging.getLogger(__name__)


def extract_docx(file_path: str) -> str | None:
    """
    Extracts text content from Word documents (.docx, .doc) using python-docx.
    For legacy .doc files, falls back to attempting extraction via docx2txt if available.
    """
    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "docx":
        return _extract_docx(file_path)
    elif ext == "doc":
        return _extract_doc(file_path)

    logger.warning(f"extract_docx called with unsupported extension: {ext}")
    return None


def _extract_docx(file_path: str) -> str | None:
    """Extract text from modern .docx format using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n\n".join(paragraphs) if paragraphs else ""
    except Exception as e:
        logger.error(f"python-docx extraction failed for {file_path}: {e}")
        return None


def _extract_doc(file_path: str) -> str | None:
    """Extract text from legacy .doc format using docx2txt or antiword fallback."""
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        return text.strip() if text and text.strip() else ""
    except ImportError:
        logger.warning("docx2txt not available for .doc extraction; text will be skipped.")
        return None
    except Exception as e:
        logger.error(f"docx2txt extraction failed for {file_path}: {e}")
        return None
